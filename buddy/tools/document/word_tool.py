from __future__ import annotations

import asyncio
import base64
import html.parser
import io
import os
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Tuple

from buddy.prompts.word_prompts import WORD_TOOL_PROMPT
from buddy.tools.document.document_utils import extract_docx_to_html, search_html

_TOOL = "word"

# NOTE: extract_docx_to_html must stamp paragraph IDs as  p0, p1, p2…
#       and table IDs as  t0, t1, t2…  so they round-trip correctly
#       through _edit's section_id parser.


# ── helpers ───────────────────────────────────────────────────────────────────


def _resolve(raw: str) -> str:
    p = os.path.expanduser(os.path.expandvars(raw.strip()))
    if not os.path.isabs(p):
        p = os.path.join(os.path.expanduser("~"), p)
    return p


def _ok(**kw: Any) -> Dict[str, Any]:
    return {"STATUS": "success", "TOOL": _TOOL, **kw}


def _err(msg: str, **kw: Any) -> Dict[str, Any]:
    return {"STATUS": "failed", "TOOL": _TOOL, "ERROR": msg, **kw}


def _needs_confirm(preview: str, **kw: Any) -> Dict[str, Any]:
    # STATUS is "needs_confirmation", NOT "failed" — so the AI knows to retry
    return {
        "STATUS": "needs_confirmation",
        "TOOL": _TOOL,
        "NEEDS_CONFIRMATION": True,
        "PREVIEW": preview,
        "NOTE": "Call again with confirmed=true after user approves.",
        **kw,
    }


# ── paragraph helpers ─────────────────────────────────────────────────────────


def _clear_paragraph(para: Any) -> None:
    """
    Remove all runs/content from a paragraph while preserving its
    paragraph-level properties (pPr: style, indentation, spacing, etc.).
    """
    from lxml import etree

    p_elem = para._p
    to_remove = [child for child in p_elem if etree.QName(child).localname != "pPr"]
    for child in to_remove:
        p_elem.remove(child)


def _patch_runs(para: Any, old_text: str, new_text: str, warnings: List[str]) -> bool:
    """
    Replace old_text → new_text inside a paragraph, preserving run formatting
    when possible.  Falls back to a flat replacement (losing per-run styles)
    only when the search text spans multiple runs.
    Returns True if the replacement was made.
    """
    # Fast path: text lives in a single run → format preserved
    for run in para.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text, 1)
            return True
    # Slow path: text spans runs → flatten paragraph
    full = para.text
    if old_text not in full:
        return False
    warnings.append(
        f"'{old_text}' spans multiple formatted runs; "
        "run-level formatting in that paragraph was simplified."
    )
    _clear_paragraph(para)
    para.add_run(full.replace(old_text, new_text, 1))
    return True


# ── HTML fragment parser ───────────────────────────────────────────────────────


class _HTMLInterceptor(html.parser.HTMLParser):
    """
    Converts a small HTML fragment into python-docx runs on a single paragraph.
    Handles:  <b>/<strong>, <i>/<em>, <u>, <br>, <img src="…|data:…">
    Uses a proper stack for nested tags so <b>outer <b>inner</b> still bold</b>
    works correctly.
    """

    def __init__(self, para: Any, warnings: List[str]) -> None:
        super().__init__()
        self.para = para
        self.warnings = warnings
        # Each stack entry is the count of open tags of that type
        self._bold_depth = 0
        self._italic_depth = 0
        self._underline_depth = 0

    # -- tag tracking ---------------------------------------------------------

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag in ("b", "strong"):
            self._bold_depth += 1
        elif tag in ("i", "em"):
            self._italic_depth += 1
        elif tag == "u":
            self._underline_depth += 1
        elif tag == "br":
            self.para.add_run().add_break()
        elif tag == "img":
            self._handle_img(dict(attrs))

    def handle_endtag(self, tag: str) -> None:
        if tag in ("b", "strong"):
            self._bold_depth = max(0, self._bold_depth - 1)
        elif tag in ("i", "em"):
            self._italic_depth = max(0, self._italic_depth - 1)
        elif tag == "u":
            self._underline_depth = max(0, self._underline_depth - 1)

    def handle_data(self, data: str) -> None:
        if not data:
            return
        run = self.para.add_run(data)
        run.bold = self._bold_depth > 0
        run.italic = self._italic_depth > 0
        run.underline = self._underline_depth > 0

    # -- image ----------------------------------------------------------------

    def _handle_img(self, attrs: Dict[str, str]) -> None:
        src = attrs.get("src", "")
        if not src:
            return
        try:
            run = self.para.add_run()
            if src.startswith("data:image"):
                _, b64 = src.split(",", 1)
                run.add_picture(io.BytesIO(base64.b64decode(b64)))
            else:
                if src.startswith("~"):
                    src = os.path.expanduser(src)
                run.add_picture(src)
        except Exception as exc:
            self.warnings.append(f"Image could not be embedded ({src[:60]}…): {exc}")


def _parse_and_apply(html_str: str, para: Any, warnings: List[str]) -> None:
    """Clear a paragraph and fill it from an HTML fragment."""
    _clear_paragraph(para)
    html_str = html_str.strip()
    # Strip a single wrapping <p>…</p> if present so callers can pass either form
    html_str = re.sub(r"(?i)^<p[^>]*>|</p>$", "", html_str).strip()
    _HTMLInterceptor(para, warnings).feed(html_str)


# ── section-id helpers ────────────────────────────────────────────────────────


def _resolve_target(doc: Any, section_id: str) -> Tuple[Optional[Any], Optional[str]]:
    """
    Parse section_id of the form  p<N>  or  t<N>  and return
    (element, type_char) or (None, None) on any error.
    """
    if not section_id or len(section_id) < 2:
        return None, None
    t_type = section_id[0].lower()
    try:
        idx = int(section_id[1:])
    except ValueError:
        return None, None
    if t_type == "p" and 0 <= idx < len(doc.paragraphs):
        return doc.paragraphs[idx], "p"
    if t_type == "t" and 0 <= idx < len(doc.tables):
        return doc.tables[idx], "t"
    return None, None


# ── tool ──────────────────────────────────────────────────────────────────────


class WordTool:
    tool_name = _TOOL
    version = "1.1.0"

    # ── info ──────────────────────────────────────────────────────────────────

    def get_info(self) -> Dict[str, Any]:
        return {
            "name": self.tool_name,
            "version": self.version,
            "description": (
                "WHEN: creating or editing .docx documents, or converting other file"
                " formats into .docx.\n\nFUNCTIONS:\n  create(path, content)       —"
                " HTML+CSS string → .docx; saves .html source alongside for future"
                " edits\n  read(path, search?)         — returns full HTML source with"
                " section IDs stamped (id='s1','s2',...); search= returns only matching"
                " sections\n  edit(path, edits[])         — patch HTML and re-render:"
                " {section_id+new} replace, {old+new} text patch,"
                " {op:add_after/add_before/add_end+new} insert, {op:remove+section_id}"
                " delete\n  convert(source, target)     — convert .pdf/.html/.md/.txt →"
                " .docx\n  export(path, target)        — .docx → .pdf via"
                " LibreOffice\n\nCHAIN: always call read before edit to get current"
                " section IDs — IDs change after every edit. convert output path feeds"
                " read/edit for further changes.\nNOT: .xlsx → excel | .pdf creation →"
                " pdf | plain file reads → fs_read | plain file writes → fs_write"
            ),
            "prompt": WORD_TOOL_PROMPT,
        }

    # ── dispatch ──────────────────────────────────────────────────────────────

    async def execute(
        self,
        function: str,
        arguments: Dict[str, Any],
        on_progress: Optional[Callable] = None,
        goal: str = "",
        brain: Any = None,
        **_: Any,
    ) -> Dict[str, Any]:
        fn = str(function or "").strip().lower()
        handlers: Dict[str, Any] = {
            "create": self._create,
            "read": self._read,
            "edit": self._edit,
            "convert": self._convert,
            "export": self._export,
        }
        if fn not in handlers:
            return _err(
                f"Unknown function: {function!r}. Must be one of: {', '.join(handlers)}"
            )
        # All handlers are synchronous and may do blocking I/O —
        # run them in a thread pool to avoid blocking the event loop.
        return await asyncio.to_thread(handlers[fn], arguments)

    # ── create ────────────────────────────────────────────────────────────────

    def _create(self, args: Dict[str, Any]) -> Dict[str, Any]:
        raw = str(args.get("path") or "").strip()
        if not raw:
            return _err(
                "word.create: 'path' is required — provide an absolute .docx path"
            )
        if not raw.lower().endswith(".docx"):
            return _err("word.create: path must end in .docx")

        content = str(args.get("content") or "").strip()
        if not content:
            return _err(
                "word.create: 'content' is required — provide an HTML+CSS string"
            )

        path = _resolve(raw)
        p = Path(path)

        if p.exists() and not args.get("confirmed"):
            return _needs_confirm(
                f"File already exists: {path}\nPass confirmed=true to overwrite it."
            )

        try:
            from docx import Document
            from htmldocx import HtmlToDocx
        except ImportError as exc:
            return _err(
                f"Missing dependency: {exc}. Run: pip install python-docx htmldocx"
            )

        try:
            p.parent.mkdir(parents=True, exist_ok=True)
            doc = Document()
            converter = HtmlToDocx()
            converter.add_html_to_document(content, doc)  # ← correct API
            doc.save(str(p))
            return _ok(PATH=str(p), SIZE_BYTES=p.stat().st_size)
        except PermissionError:
            return _err(f"Permission denied: {path}")
        except Exception as exc:
            return _err(f"Failed to create document: {exc}")

    # ── read ──────────────────────────────────────────────────────────────────

    def _read(self, args: Dict[str, Any]) -> Dict[str, Any]:
        raw = str(args.get("path") or "").strip()
        if not raw:
            return _err(
                "word.read: 'path' is required — provide an absolute .docx path"
            )

        path = _resolve(raw)
        p = Path(path)
        if not p.exists():
            return _err(f"word.read: file not found: {path}")

        search = str(args.get("search") or "").strip() or None

        try:
            html = extract_docx_to_html(path)
        except ImportError as exc:
            return _err(str(exc))
        except Exception as exc:
            return _err(f"Failed to extract document content: {exc}")

        if search:
            html = search_html(html, search)
            if not html:
                return _ok(
                    PATH=str(p),
                    SEARCH=search,
                    HTML="",
                    NOTE="No sections matched the search query.",
                )

        return _ok(PATH=str(p), HTML=html)

    # ── edit ──────────────────────────────────────────────────────────────────

    def _edit(self, args: Dict[str, Any]) -> Dict[str, Any]:
        raw = str(args.get("path") or "").strip()
        if not raw:
            return _err(
                "word.edit: 'path' is required — provide an absolute .docx path"
            )

        edits = args.get("edits")
        if not edits or not isinstance(edits, list):
            return _err("word.edit: 'edits' must be a non-empty list")

        path = _resolve(raw)
        p = Path(path)
        if not p.exists():
            return _err(f"word.edit: file not found: {path}")

        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_BREAK
            from docx.enum.section import WD_ORIENT
            from docx.text.paragraph import Paragraph as DocxParagraph
            from lxml.etree import Element
        except ImportError as exc:
            return _err(f"Missing dependency: {exc}. Run: pip install python-docx lxml")

        # ── backup before touching the file ───────────────────────────────────
        backup = path + ".bak"
        try:
            shutil.copy2(path, backup)
        except Exception as exc:
            return _err(f"Could not create backup before editing: {exc}")

        try:
            doc = Document(path)
        except Exception as exc:
            _safe_remove(backup)
            return _err(f"Failed to open document: {exc}")

        applied: List[Dict] = []
        failed: List[Dict] = []
        warnings: List[str] = []

        _MARGINS = {
            "narrow": (0.5, 0.5, 0.5, 0.5),
            "normal": (1.0, 1.0, 1.0, 1.0),
            "wide": (1.0, 1.0, 2.0, 2.0),
        }
        _VALID_OPS = {
            "replace",
            "patch",
            "remove",
            "add_after",
            "add_before",
            "add_end",
            "add_page_break",
            "set_page_setup",
        }

        for edit in edits:
            op = str(edit.get("op") or "").strip()
            section_id = str(edit.get("section_id") or "")

            if not op:
                failed.append({"edit": edit, "error": "Missing 'op' field"})
                continue
            if op not in _VALID_OPS:
                failed.append({
                    "op": op,
                    "error": f"Unknown op. Valid: {', '.join(sorted(_VALID_OPS))}",
                })
                continue

            try:
                # ── ops that don't require section_id ─────────────────────────

                if op == "set_page_setup":
                    margin = edit.get("margin")
                    orient = edit.get("orientation")
                    for section in doc.sections:
                        if margin in _MARGINS:
                            t, b, l, r = _MARGINS[margin]
                            section.top_margin = Inches(t)
                            section.bottom_margin = Inches(b)
                            section.left_margin = Inches(l)
                            section.right_margin = Inches(r)
                        if orient == "landscape":
                            section.orientation = WD_ORIENT.LANDSCAPE
                            w, h = section.page_width, section.page_height
                            if w is not None and h is not None and w < h:
                                section.page_width, section.page_height = h, w
                        elif orient == "portrait":
                            section.orientation = WD_ORIENT.PORTRAIT
                            w, h = section.page_width, section.page_height
                            if w is not None and h is not None and w > h:
                                section.page_width, section.page_height = h, w
                    applied.append({"op": op})
                    continue

                if op == "add_end":
                    new_p = doc.add_paragraph()
                    _apply_style(new_p, edit.get("style"), warnings)
                    _parse_and_apply(edit.get("new", ""), new_p, warnings)
                    applied.append({"op": op})
                    continue

                # ── ops that require section_id ───────────────────────────────

                if not section_id:
                    failed.append(
                        {"op": op, "error": "section_id is required for this op"}
                    )
                    continue

                target, t_type = _resolve_target(doc, section_id)
                if target is None:
                    failed.append({
                        "op": op,
                        "section_id": section_id,
                        "error": (
                            "section_id not found. IDs change after every edit — call"
                            " read again to refresh them."
                        ),
                    })
                    continue

                # replace ─────────────────────────────────────────────────────
                if op == "replace":
                    if t_type != "p":
                        failed.append({
                            "op": op,
                            "section_id": section_id,
                            "error": (
                                "replace is only supported on paragraphs (p…). "
                                "To edit a table cell, target its paragraph."
                            ),
                        })
                        continue
                    _parse_and_apply(edit.get("new", ""), target, warnings)
                    applied.append({"op": op, "section_id": section_id})

                # patch ───────────────────────────────────────────────────────
                elif op == "patch":
                    old_text = str(edit.get("old") or "")
                    new_text = str(edit.get("new") or "")
                    if not old_text:
                        failed.append({
                            "op": op,
                            "section_id": section_id,
                            "error": "'old' text is required for patch",
                        })
                        continue
                    if t_type != "p":
                        failed.append({
                            "op": op,
                            "section_id": section_id,
                            "error": "patch is only supported on paragraphs",
                        })
                        continue
                    if not _patch_runs(target, old_text, new_text, warnings):
                        failed.append({
                            "op": op,
                            "section_id": section_id,
                            "error": f"'{old_text}' not found in paragraph",
                        })
                        continue
                    applied.append({"op": op, "section_id": section_id})

                # remove ──────────────────────────────────────────────────────
                elif op == "remove":
                    element = target._element
                    parent = element.getparent()
                    if parent is None:
                        failed.append({
                            "op": op,
                            "section_id": section_id,
                            "error": "Element has no parent; cannot remove.",
                        })
                        continue
                    parent.remove(element)
                    applied.append({"op": op, "section_id": section_id})

                # add_after / add_before ───────────────────────────────────────
                elif op in ("add_after", "add_before"):
                    if t_type != "p":
                        failed.append({
                            "op": op,
                            "section_id": section_id,
                            "error": f"{op} is only supported on paragraphs",
                        })
                        continue
                    element = target._element
                    new_element = Element(element.tag)
                    if op == "add_after":
                        element.addnext(new_element)
                    else:
                        element.addprevious(new_element)
                    new_p = DocxParagraph(new_element, target._parent)
                    _apply_style(new_p, edit.get("style"), warnings)
                    _parse_and_apply(edit.get("new", ""), new_p, warnings)
                    applied.append({"op": op, "section_id": section_id})

                # add_page_break ──────────────────────────────────────────────
                elif op == "add_page_break":
                    if t_type != "p":
                        failed.append({
                            "op": op,
                            "section_id": section_id,
                            "error": "Page break can only be added to paragraphs",
                        })
                        continue
                    target.add_run().add_break(WD_BREAK.PAGE)
                    applied.append({"op": op, "section_id": section_id})

            except Exception as exc:
                failed.append({"edit": edit, "error": str(exc)})

        # ── atomic save: restore backup on any failure ─────────────────────────
        if failed:
            try:
                shutil.copy2(backup, path)  # restore original
            except Exception:
                pass
            _safe_remove(backup)
            out: Dict[str, Any] = {
                "STATUS": "failed",
                "TOOL": _TOOL,
                "PATH": str(p),
                "EDITS_APPLIED": 0,
                "EDITS_FAILED": len(failed),
                "FAILED": failed,
                "ERROR": (
                    f"{len(failed)} edit(s) failed — no changes were saved. "
                    "Fix the errors and retry."
                ),
            }
            if warnings:
                out["WARNINGS"] = warnings
            return out

        try:
            doc.save(path)
        except PermissionError:
            _restore_backup(backup, path)
            return _err(f"Permission denied when saving: {path}")
        except Exception as exc:
            _restore_backup(backup, path)
            return _err(f"Save failed (original restored from backup): {exc}")
        finally:
            _safe_remove(backup)

        out = {
            "STATUS": "success",
            "TOOL": _TOOL,
            "PATH": str(p),
            "EDITS_APPLIED": len(applied),
            "APPLIED": applied,
        }
        if warnings:
            out["WARNINGS"] = warnings
        return out

    # ── convert ───────────────────────────────────────────────────────────────

    def _convert(self, args: Dict[str, Any]) -> Dict[str, Any]:
        source_raw = str(args.get("source") or "").strip()
        target_raw = str(args.get("target") or "").strip()
        if not source_raw:
            return _err("word.convert: 'source' is required")
        if not target_raw:
            return _err("word.convert: 'target' is required")
        if not target_raw.lower().endswith(".docx"):
            return _err("word.convert: 'target' must end in .docx")

        source_path = _resolve(source_raw)
        target_path = _resolve(target_raw)
        sp = Path(source_path)
        if not sp.exists():
            return _err(f"word.convert: source file not found: {source_path}")

        ext = sp.suffix.lower()
        supported = {".html", ".htm", ".md", ".txt", ".pdf"}
        if ext not in supported:
            return _err(
                f"word.convert: unsupported source format '{ext}'. "
                f"Supported: {', '.join(sorted(supported))}"
            )

        try:
            Path(target_path).parent.mkdir(parents=True, exist_ok=True)

            if ext in (".html", ".htm"):
                from docx import Document
                from htmldocx import HtmlToDocx

                doc = Document()
                HtmlToDocx().add_html_to_document(
                    sp.read_text(encoding="utf-8", errors="replace"), doc
                )
                doc.save(target_path)

            elif ext == ".md":
                try:
                    import markdown as md_lib
                except ImportError:
                    return _err(
                        "Markdown conversion requires 'markdown'. Run: pip install"
                        " markdown"
                    )
                from docx import Document
                from htmldocx import HtmlToDocx

                html_content = md_lib.markdown(
                    sp.read_text(encoding="utf-8", errors="replace")
                )
                doc = Document()
                HtmlToDocx().add_html_to_document(html_content, doc)
                doc.save(target_path)

            elif ext == ".txt":
                from docx import Document

                doc = Document()
                for line in sp.read_text(
                    encoding="utf-8", errors="replace"
                ).splitlines():
                    doc.add_paragraph(line)
                doc.save(target_path)

            elif ext == ".pdf":
                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "docx",
                        "--outdir",
                        str(Path(target_path).parent),
                        source_path,
                    ],
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
                if result.returncode != 0:
                    return _err(
                        f"LibreOffice conversion failed: {result.stderr.strip()}"
                    )
                # LibreOffice outputs to <source_stem>.docx; rename if needed
                lo_out = Path(target_path).parent / sp.with_suffix(".docx").name
                if lo_out != Path(target_path):
                    lo_out.rename(target_path)

            tp = Path(target_path)
            return _ok(
                SOURCE=source_path, PATH=target_path, SIZE_BYTES=tp.stat().st_size
            )

        except FileNotFoundError:
            return _err("LibreOffice not found. Install it to convert PDF files.")
        except subprocess.TimeoutExpired:
            return _err("Conversion timed out after 60 seconds.")
        except PermissionError:
            return _err(f"Permission denied: {target_path}")
        except Exception as exc:
            return _err(f"Conversion failed: {exc}")

    # ── export ────────────────────────────────────────────────────────────────

    def _export(self, args: Dict[str, Any]) -> Dict[str, Any]:
        raw = str(args.get("path") or "").strip()
        target = str(args.get("target") or "").strip()
        if not raw:
            return _err("word.export: 'path' is required — absolute .docx path")
        if not target:
            return _err("word.export: 'target' is required — absolute .pdf path")
        if not target.lower().endswith(".pdf"):
            return _err("word.export: 'target' must end in .pdf")

        path = _resolve(raw)
        target_path = _resolve(target)
        p = Path(path)
        if not p.exists():
            return _err(f"word.export: file not found: {path}")

        try:
            Path(target_path).parent.mkdir(parents=True, exist_ok=True)
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(Path(target_path).parent),
                    path,
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                return _err(f"LibreOffice PDF export failed: {result.stderr.strip()}")
            lo_out = Path(target_path).parent / p.with_suffix(".pdf").name
            if lo_out != Path(target_path):
                lo_out.rename(target_path)
            tp = Path(target_path)
            return _ok(SOURCE=path, PATH=target_path, SIZE_BYTES=tp.stat().st_size)
        except FileNotFoundError:
            return _err("LibreOffice not found. Install it to enable PDF export.")
        except subprocess.TimeoutExpired:
            return _err("PDF export timed out after 60 seconds.")
        except PermissionError:
            return _err(f"Permission denied: {target_path}")
        except Exception as exc:
            return _err(f"Export failed: {exc}")


# ── private utilities ─────────────────────────────────────────────────────────


def _apply_style(para: Any, style: Optional[str], warnings: List[str]) -> None:
    if not style:
        return
    try:
        para.style = style
    except Exception:
        warnings.append(
            f"Style '{style}' not found in this document; default style used."
        )


def _safe_remove(path: str) -> None:
    try:
        os.remove(path)
    except Exception:
        pass


def _restore_backup(backup: str, path: str) -> None:
    try:
        shutil.copy2(backup, path)
    except Exception:
        pass
    _safe_remove(backup)


# ── registry ──────────────────────────────────────────────────────────────────

TOOL_NAME = "word"
TOOL_CLASS = WordTool


def get_tool() -> WordTool:
    return WordTool()
